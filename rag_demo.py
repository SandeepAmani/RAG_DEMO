# =============================================================================
# RAG DEMO — Retrieval-Augmented Generation for Study Purposes
# =============================================================================
#
# WHAT IS RAG?
# ------------
# RAG stands for Retrieval-Augmented Generation.
# It combines two ideas:
#   1. RETRIEVAL  — search a knowledge base to find relevant text chunks
#   2. GENERATION — send those chunks to an LLM so it can answer grounded
#                   in actual content (not just its training memory)
#
# WHY USE RAG?
# ------------
# LLMs like GPT or LLaMA are trained on fixed data. They can hallucinate
# (make things up) or not know recent/private facts. RAG fixes this by
# giving the LLM fresh, relevant context at query time.
#
# PIPELINE (end-to-end flow):
# ---------------------------
#   Documents (text OR images)
#        ↓
#   Chunking / Image encoding (split text into pieces; encode images to base64)
#        ↓
#   Embedding (convert each chunk/image to a vector using CLIP multimodal model)
#        ↓
#   Vector Store (store vectors in memory for fast search)
#        ↓
#   User asks a question
#        ↓
#   Question → Embedding → Search vector store → Top-K chunks/images retrieved
#        ↓
#   Prompt = "Use this context: <chunks + images> \n Answer: <question>"
#        ↓
#   LLM (Groq / LLaMA3-Vision) generates the final answer
#
# COMPONENTS USED IN THIS SCRIPT:
# --------------------------------
#   • sentence-transformers  — HuggingFace library; now using CLIP for multimodal embeddings
#   • numpy                  — for vector math (cosine similarity search)
#   • groq (Python SDK)      — to call the Groq LLM API (vision-capable model)
#   • os / getpass           — to read the API key from environment variables
#   • base64                 — to encode images for the Groq vision API
#   • pdf2image + pytesseract — OCR fallback for scanned/image-based PDFs
#   • PIL (Pillow)           — to open and pre-process image files
#
# HOW TO RUN:
# -----------
#   1. Install dependencies:
#        pip install sentence-transformers numpy groq pypdf python-docx
#        pip install Pillow pdf2image pytesseract
#        (Windows: also install Tesseract OCR from https://github.com/UB-Mannheim/tesseract/wiki)
#        (Windows: also install Poppler from https://github.com/oschwartz10612/poppler-windows/releases)
#
#   2. Set your Groq API key (get a free one at https://console.groq.com):
#        Windows PowerShell:   $env:GROQ_API_KEY = "your_key_here"
#        Windows CMD:          set GROQ_API_KEY=your_key_here
#        macOS/Linux bash:     export GROQ_API_KEY=your_key_here
#
#   3. Run the script:
#        python rag_demo.py
# =============================================================================


# --- STANDARD LIBRARY IMPORTS ------------------------------------------------
import os        # Used to read environment variables (e.g., GROQ_API_KEY)
import sys       # Used to exit the script cleanly on errors
import textwrap  # Used to wrap long text output neatly in the terminal
import base64    # [NEW] Used to encode image bytes to base64 for the Groq vision API
from pathlib import Path  # Used to locate the .env file relative to this script

# --- THIRD-PARTY IMPORTS -----------------------------------------------------
# These must be installed via: pip install sentence-transformers numpy groq pypdf python-docx

try:
    import numpy as np
    # numpy is a numerical computing library.
    # We use it here to:
    #   - Store embedding vectors as arrays
    #   - Compute cosine similarity (how similar two vectors are)
except ImportError:
    print("ERROR: 'numpy' is not installed. Run: pip install numpy")
    sys.exit(1)

try:
    from sentence_transformers import SentenceTransformer
    # SentenceTransformer loads a pre-trained HuggingFace embedding model.
    # It converts a string of text → a list of numbers (a "vector").
    # Similar texts produce vectors that are close together in vector space.
    #
    # [CHANGED] The model is now 'clip-ViT-B-32' (multimodal) instead of 'all-MiniLM-L6-v2' (text only).
    # CLIP (Contrastive Language-Image Pretraining) by OpenAI maps BOTH text and images
    # into the same 512-dimensional vector space, so text queries can retrieve image chunks
    # and vice versa. The old text-only model is commented out below at load_embedding_model().
except ImportError:
    print("ERROR: 'sentence-transformers' is not installed.")
    print("Run: pip install sentence-transformers")
    sys.exit(1)

try:
    from groq import Groq
    # Groq is the Python SDK for calling the Groq LLM API.
    # Groq provides free, very fast inference for open-source models like LLaMA3.
    # We use it to generate the final natural-language answer.
    #
    # [CHANGED] We now use 'meta-llama/llama-4-scout-17b-16e-instruct' instead of 'llama-3.1-8b-instant'.
    # The vision model can receive base64-encoded images alongside text in the same API call,
    # enabling true multimodal Q&A over image content.
except ImportError:
    print("ERROR: 'groq' is not installed. Run: pip install groq")
    sys.exit(1)

try:
    from pypdf import PdfReader
    # pypdf reads PDF files page by page and extracts their text content.
    # PdfReader opens the file; .pages gives a list of page objects;
    # page.extract_text() returns the text on that page as a plain string.
    # NOTE: For scanned/image PDFs, extract_text() returns "". We now fall back to OCR.
except ImportError:
    print("ERROR: 'pypdf' is not installed. Run: pip install pypdf")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    # python-docx reads .docx (Word) files.
    # Document() opens the file; .paragraphs gives a list of paragraph objects;
    # paragraph.text returns the text of that paragraph as a plain string.
except ImportError:
    print("ERROR: 'python-docx' is not installed. Run: pip install python-docx")
    sys.exit(1)

# [NEW] Pillow — image loading and preprocessing
try:
    from PIL import Image
    # Pillow (PIL fork) opens common image formats: PNG, JPG, GIF, WebP, etc.
    # We use it to open image files and convert them to RGB before encoding.
except ImportError:
    print("WARNING: 'Pillow' is not installed. Image file support will be disabled.")
    print("Run: pip install Pillow")
    Image = None  # graceful degradation: image loading functions will skip

# [NEW] pymupdf (fitz) — renders PDF pages to pixel images for vision LLM.
# Self-contained: no Poppler or Tesseract needed.
# Install: pip install pymupdf
try:
    import fitz  # pymupdf
    _FITZ_AVAILABLE = True
except ImportError:
    print("WARNING: 'pymupdf' not installed. Scanned PDF pages will be skipped.")
    print("Run: pip install pymupdf")
    _FITZ_AVAILABLE = False


# =============================================================================
# STEP 0 — DOCUMENT LOADER (load text from PDF, Word, or Image files)
# =============================================================================
#
# [CHANGED] This section now supports three document types:
#   - PDF files   (.pdf)  — text extraction; scanned pages rendered to images via pymupdf
#   - Word files  (.docx) — paragraph text extraction (unchanged)
#   - Image files (.png, .jpg, .jpeg, .gif, .webp) — base64 encoded for vision LLM
#
# HOW IT WORKS:
#   - load_pdf(path)   → returns (text_str, [image_dicts]) — text pages as text,
#                        scanned/image pages rendered by pymupdf as image dicts
#   - load_docx(path)  → reads every paragraph (unchanged from before)
#   - load_image(path) → reads image file, returns a special dict with base64 data
#   - load_documents_from_files(paths) → auto-detects file type, returns a list of
#                                        text strings OR image dicts

def load_pdf(file_path: str) -> tuple:
    """
    Extract content from a PDF file.

    Strategy:
      - For each page, first try pypdf text extraction (instant, no dependencies).
      - If a page has no extractable text (scanned page, image-only page, diagram),
        render it to a PNG using pymupdf (fitz) and return it as an image dict.
        The Groq vision LLM will then read the image directly — no OCR needed.

    Why pymupdf instead of pdf2image + pytesseract?
      - pymupdf is self-contained (pip install pymupdf) — no Poppler or Tesseract.
      - Rendering to image + sending to vision LLM is more accurate than OCR for
        diagrams, tables, and mixed-layout documents.

    Parameters:
        file_path (str): absolute or relative path to the .pdf file

    Returns:
        text      (str)  : joined text from all text-layer pages
        img_pages (list) : image dicts (type='image') for pages with no text
    """
    print(f"  [PDF] Reading: {file_path}")
    reader = PdfReader(file_path)
    pages_text = []
    img_pages = []

    # Open same file with pymupdf for rendering image-based pages
    fitz_doc = fitz.open(file_path) if _FITZ_AVAILABLE else None

    for i, page in enumerate(reader.pages):
        # extract_text() returns the raw text layer — empty string for scanned pages
        text = page.extract_text()

        if text and text.strip():
            # Page has a text layer — use it directly (same as before)
            pages_text.append(text.strip())
        else:
            # Page is image-based (scanned, diagram, chart, etc.)
            if fitz_doc:
                # [NEW] Render the page to a PNG image using pymupdf
                # Matrix(2, 2) = 2× zoom = 144 DPI — good quality without huge file size
                print(f"  [PDF] Page {i+1} has no text — rendering as image for vision LLM...")
                try:
                    fitz_page = fitz_doc[i]
                    mat = fitz.Matrix(2, 2)
                    pix = fitz_page.get_pixmap(matrix=mat, alpha=False)
                    png_bytes = pix.tobytes("png")
                    b64_data = base64.b64encode(png_bytes).decode("utf-8")
                    img_pages.append({
                        "type": "image",
                        "path": f"{Path(file_path).name} — page {i+1}",
                        "mime_type": "image/png",
                        "b64_data": b64_data,
                    })
                    print(f"  [PDF] Page {i+1} rendered: {len(png_bytes):,} bytes")
                except Exception as e:
                    print(f"  [PDF] Failed to render page {i+1}: {e}")
            else:
                print(f"  [PDF] Page {i+1} has no text. Install pymupdf to read it: pip install pymupdf")

    if fitz_doc:
        fitz_doc.close()

    combined = "\n\n".join(pages_text)
    print(f"  [PDF] {len(reader.pages)} pages: {len(pages_text)} text pages, {len(img_pages)} image pages.")
    return combined, img_pages


def load_docx(file_path: str) -> str:
    """
    Extract all text from a Word (.docx) file.
    [UNCHANGED] This function is exactly as before — Word docs are text-based.

    Parameters:
        file_path (str): absolute or relative path to the .docx file

    Returns:
        text (str): all paragraph text joined together, or empty string on failure
    """
    print(f"  [DOCX] Reading: {file_path}")
    doc = DocxDocument(file_path)
    # doc.paragraphs is a list — one Paragraph object per paragraph block.
    # paragraph.text gives the plain text (no formatting codes).
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    combined = "\n\n".join(paragraphs)
    print(f"  [DOCX] Extracted {len(paragraphs)} paragraphs, "
          f"{len(combined):,} characters of text.")
    return combined


def load_image(file_path: str) -> dict | None:
    """
    [NEW] Load an image file and encode it as base64 for the Groq vision API.

    Why base64? The Groq vision API (like most multimodal LLM APIs) accepts images
    as base64-encoded strings inside the message content array. This lets us send
    image data inline without needing a public URL.

    The returned dict has a special 'type': 'image' key so downstream code can
    distinguish image chunks from text chunks and handle them differently when
    building the LLM prompt.

    Parameters:
        file_path (str): absolute path to a .png, .jpg, .jpeg, .gif, or .webp file

    Returns:
        dict with keys:
            'type'      : 'image'
            'path'      : original file path (for display)
            'mime_type' : e.g. 'image/jpeg'
            'b64_data'  : base64-encoded image bytes as a string
        or None if Pillow is not installed or the file cannot be opened
    """
    if Image is None:
        print(f"  [IMAGE] Skipping {file_path} — Pillow not installed.")
        return None

    suffix = Path(file_path).suffix.lower()
    mime_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    mime_type = mime_map.get(suffix, "image/jpeg")

    print(f"  [IMAGE] Loading: {file_path}")
    try:
        # Open the image with Pillow and convert to RGB
        # (Some images are RGBA or palette mode — RGB is universally supported)
        img = Image.open(file_path).convert("RGB")

        # Save the image to an in-memory byte buffer as JPEG
        import io
        buf = io.BytesIO()
        img.save(buf, format="JPEG" if mime_type != "image/png" else "PNG")
        img_bytes = buf.getvalue()

        # base64.b64encode() converts raw bytes → base64 bytes
        # .decode("utf-8") turns those bytes into a plain Python string
        b64_data = base64.b64encode(img_bytes).decode("utf-8")

        print(f"  [IMAGE] Encoded {len(img_bytes):,} bytes → {len(b64_data):,} base64 chars")
        return {
            "type": "image",
            "path": file_path,
            "mime_type": mime_type,
            "b64_data": b64_data,
        }
    except Exception as e:
        print(f"  [IMAGE] Failed to load {file_path}: {e}")
        return None


def load_documents_from_files(file_paths: list) -> list:
    """
    Load text OR image data from a list of file paths.

    [CHANGED] Previously only supported .pdf and .docx.
    Now also handles image files (.png, .jpg, .jpeg, .gif, .webp) by calling
    load_image() which returns a dict instead of a string. The returned list
    can contain a mix of strings (text docs) and dicts (image docs).

    Parameters:
        file_paths (list of str): paths to .pdf, .docx, or image files

    Returns:
        documents (list): mix of str (text) and dict (image) per successfully loaded file
    """
    documents = []
    for path in file_paths:
        # Normalise: strip 'file:///' prefix if the user pastes a browser URL
        # e.g. "file:///C:/Users/..." → "C:/Users/..."
        if path.startswith("file:///"):
            path = path[8:]            # remove "file:///"
        elif path.startswith("file://"):
            path = path[7:]            # remove "file://"

        path = path.replace("/", "\\") # normalise slashes on Windows

        if not Path(path).exists():
            print(f"  [WARNING] File not found, skipping: {path}")
            continue

        suffix = Path(path).suffix.lower()  # e.g. ".pdf", ".docx", ".png"

        if suffix == ".pdf":
            # load_pdf now returns (text, img_pages) — handle both
            text, img_pages = load_pdf(path)
            if text.strip():
                documents.append(text)
            if img_pages:
                # Scanned/image PDF pages returned as image dicts for the vision LLM
                documents.extend(img_pages)
                print(f"  [PDF] Added {len(img_pages)} image page(s) from: {path}")
            if not text.strip() and not img_pages:
                print(f"  [WARNING] No content extracted from: {path}")

        elif suffix in (".docx", ".doc"):
            text = load_docx(path)
            if text.strip():
                documents.append(text)
            else:
                print(f"  [WARNING] No text extracted from: {path}")

        # [NEW] Image file support — these were previously logged as unsupported and skipped
        elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
            img_doc = load_image(path)
            if img_doc:
                documents.append(img_doc)  # append the image dict (not a string)

        else:
            print(f"  [WARNING] Unsupported file type '{suffix}', skipping: {path}")
            continue

    return documents


# =============================================================================
# STEP 1 — KNOWLEDGE BASE (Hardcoded study documents about RAG)
# =============================================================================
#
# In a real production RAG system, these documents would come from PDFs,
# databases, web pages, or uploaded files. Here we hardcode them directly
# in the script so you can study without any external files.
#
# Each string in this list is one "document" — a self-contained piece of text.
# The RAG system will split these into smaller chunks, embed them, and search
# through them to answer your questions.
DOCUMENTS = []
# DOCUMENTS = [
#     """
#     RAG stands for Retrieval-Augmented Generation. It is a technique in AI
#     that combines information retrieval with text generation. Instead of
#     relying solely on what an LLM learned during training, RAG allows the model
#     to look up relevant information from an external knowledge base at query
#     time. This makes responses more accurate, factual, and up-to-date.
#     """,

#     """
#     The RAG pipeline has three main stages: indexing, retrieval, and generation.
#     During indexing, documents are split into chunks, and each chunk is converted
#     into a vector embedding using an embedding model. During retrieval, the user's
#     question is also embedded, and the most similar document chunks are found using
#     cosine similarity or other vector search methods. During generation, the
#     retrieved chunks are passed to the LLM as context so it can generate a
#     grounded answer.
#     """,

#     """
#     An embedding is a numerical representation of text — a list of floating-point
#     numbers called a vector. Semantically similar texts produce vectors that are
#     close to each other in high-dimensional space. For example, the sentences
#     "What is machine learning?" and "Define ML" would have very similar embedding
#     vectors. This property allows us to search for relevant text by comparing
#     vectors rather than doing keyword matching.
#     """,

#     """
#     Cosine similarity is the most common metric used to compare embedding vectors
#     in RAG systems. It measures the angle between two vectors — a score of 1.0
#     means the vectors point in the same direction (very similar), while a score
#     near 0 means they are unrelated. Given a user query embedding and a set of
#     document chunk embeddings, we rank the chunks by cosine similarity and return
#     the top-K most relevant chunks.
#     """,

#     """
#     Chunking is the process of splitting long documents into smaller pieces before
#     embedding them. This is important because embedding models have a maximum token
#     limit (e.g., 512 tokens). Also, smaller chunks lead to more precise retrieval —
#     you return only the paragraph that is relevant, not an entire 10-page document.
#     Common chunking strategies include: fixed-size chunking (every N characters),
#     sentence-level chunking, and semantic chunking (split on topic boundaries).
#     """,

#     """
#     The prompt sent to the LLM in a RAG system typically has three parts:
#     1. A system instruction: "You are a helpful assistant. Answer only based on
#        the provided context."
#     2. The retrieved context: the top-K document chunks found by the retriever.
#     3. The user question.
#     By instructing the LLM to answer only from context, we reduce hallucination
#     and keep responses grounded in the actual knowledge base content.
#     """,

#     """
#     Vector stores (also called vector databases) are specialized databases for
#     storing and searching embedding vectors efficiently. Examples include:
#     FAISS (Meta, open-source), ChromaDB, Pinecone, Weaviate, and Qdrant.
#     In this demo script, we implement a simple in-memory vector store using
#     a plain Python list and numpy cosine similarity — no external database needed.
#     For large-scale production use (millions of documents), a proper vector DB
#     like FAISS or Pinecone would be used instead.
#     """,

#     """
#     Groq is a company that provides extremely fast inference for open-source LLMs
#     using their custom LPU (Language Processing Unit) hardware. Their free API
#     supports models like LLaMA3 (by Meta) and Mixtral (by Mistral AI).
#     In a RAG pipeline, Groq serves as the GENERATION component — it takes the
#     retrieved context and the user question and generates a final natural-language
#     answer. The model used in this demo is llama3-8b-8192 (8 billion parameters,
#     8192 token context window).
#     """,

#     """
#     HuggingFace is an AI company and open-source platform that hosts thousands of
#     pre-trained models for NLP, computer vision, and more. The `sentence-transformers`
#     library, built on top of HuggingFace, provides easy-to-use embedding models.
#     The model used in this demo is `all-MiniLM-L6-v2` — a distilled BERT model
#     that maps sentences to a 384-dimensional embedding space. It is small (~80MB),
#     fast, runs fully offline (no API key needed), and works well for semantic search.
#     """,

#     """
#     Hallucination in LLMs refers to when the model generates plausible-sounding
#     but factually incorrect information. This happens because LLMs predict the
#     next token based on patterns in training data, not because they have verified
#     knowledge. RAG reduces hallucination by providing the model with actual source
#     text at inference time, and by instructing it to answer only from that context.
#     This is one of the core reasons RAG is widely used in production AI applications.
#     """,
# ]

# =============================================================================
# EXTERNAL FILES — Add your PDF / Word / Image documents here
# =============================================================================
#
# HOW TO USE:
#   Add the full path to your file(s) in the list below.
#   You can paste either a Windows path or a browser file:// URL — both work.
#
# EXAMPLES:
#   "C:\\Users\\sande\\Downloads\\my_document.pdf"
#   "file:///C:/Users/sande/Downloads/my_document.pdf"
#   "C:\\Users\\sande\\Documents\\notes.docx"
#   "C:\\Users\\sande\\Pictures\\diagram.png"      ← [NEW] image files now supported
#
# You can add multiple files — one path per line, separated by commas.
# Set the list to [] (empty) to use only the hardcoded study documents above.

EXTERNAL_FILES = [
    "file:///C:/Users/sande/Downloads/Learn_Coding_in_2026_12_Week_Guide.pdf",
]

# Load text from the external files and append to the DOCUMENTS list.
# This happens at import time so the rest of the pipeline sees one unified list.
if EXTERNAL_FILES:
    print("\n[FILES] Loading external documents...")
    loaded = load_documents_from_files(EXTERNAL_FILES)
    DOCUMENTS.extend(loaded)
    print(f"[FILES] Added {len(loaded)} external document(s) to knowledge base.")


# =============================================================================
# STEP 2 — CHUNKING
# =============================================================================
#
# We split each document into smaller chunks.
# In this demo, documents are already short, but we still apply chunking
# to demonstrate the concept and handle any longer documents correctly.
#
# Strategy used: fixed-size character chunking with overlap.
# - chunk_size   : max number of characters per chunk
# - chunk_overlap: how many characters to repeat between consecutive chunks
#                  Overlap helps avoid cutting a sentence mid-thought and
#                  losing context at chunk boundaries.
#
# [CHANGED] This function now handles a mix of text strings and image dicts.
# Image documents are NOT chunked — they pass through as single items.
# Only text strings are split into overlapping chunks.

def chunk_documents(documents, chunk_size=500, chunk_overlap=50):
    """
    Split text documents into overlapping chunks; pass image documents through unchanged.

    [CHANGED] Previously, `documents` was always a list of strings.
    Now it can contain image dicts (from load_image). Those are kept as-is
    and appended to the chunks list directly, because an image cannot be
    meaningfully split the way text can.

    Parameters:
        documents    (list): str (text docs) or dict (image docs)
        chunk_size   (int) : maximum characters per text chunk
        chunk_overlap(int) : characters of overlap between adjacent text chunks

    Returns:
        chunks (list): text strings OR image dicts
    """
    chunks = []

    for doc in documents:
        # [NEW] If this item is an image dict, keep it as a single chunk
        if isinstance(doc, dict) and doc.get("type") == "image":
            chunks.append(doc)   # images are passed through without splitting
            print(f"  [CHUNKING] Image kept as single chunk: {Path(doc['path']).name}")
            continue

        # Text document — apply the same fixed-size chunking as before
        doc = doc.strip()

        if len(doc) <= chunk_size:
            chunks.append(doc)
            continue

        start = 0
        while start < len(doc):
            end = min(start + chunk_size, len(doc))
            chunk = doc[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - chunk_overlap

    return chunks


# =============================================================================
# STEP 3 — EMBEDDING MODEL (HuggingFace / sentence-transformers)
# =============================================================================
#
# The embedding model converts text (or images) → vector (a list of numbers).
# We load the model ONCE at startup.
#
# [CHANGED] Model changed from 'all-MiniLM-L6-v2' to 'clip-ViT-B-32'.
#
# WHY CLIP?
# ---------
# 'all-MiniLM-L6-v2' only understands text — it has no concept of images.
# 'clip-ViT-B-32' (Contrastive Language-Image Pretraining) was trained by OpenAI
# to align text and images in the same vector space. This means:
#   - A text query like "cat diagram" produces a vector close to a cat image's vector
#   - You can retrieve image chunks using text queries and vice versa
# This is the key upgrade that makes multimodal RAG possible.
#
# CLIP dimensions: 512 (vs 384 for MiniLM)
# CLIP download size: ~600MB (vs ~80MB for MiniLM)

def load_embedding_model(model_name=None):
    """
    Load and return a SentenceTransformer embedding model.

    [CHANGED] Default model changed from 'all-MiniLM-L6-v2' to 'clip-ViT-B-32'.
    The old model name is kept as a comment below so you can see what was changed.

    Parameters:
        model_name (str): HuggingFace model identifier. Defaults to clip-ViT-B-32.

    Returns:
        model (SentenceTransformer): loaded embedding model
    """
    # [OLD] Text-only model — commented out, replaced with CLIP below
    # model_name = model_name or "all-MiniLM-L6-v2"
    # print("  (First run downloads ~80MB. Subsequent runs use local cache.)")

    # [NEW] Multimodal CLIP model — embeds both text AND images into the same space
    model_name = model_name or "clip-ViT-B-32"

    print(f"\n[EMBEDDING] Loading model '{model_name}' from HuggingFace...")
    print("  (First run downloads ~600MB for CLIP. Subsequent runs use local cache.)")

    # SentenceTransformer automatically downloads the model from HuggingFace Hub
    # and caches it at: C:\Users\<you>\.cache\huggingface\hub\  (Windows)
    model = SentenceTransformer(model_name)

    print(f"[EMBEDDING] Model loaded successfully.")
    return model


def embed_texts(model, texts):
    """
    Convert a list of text strings into embedding vectors.

    Parameters:
        model (SentenceTransformer): the loaded embedding model
        texts (list of str)        : text strings to embed

    Returns:
        embeddings (np.ndarray): 2D array of shape (N, D)
                                 N = number of texts
                                 D = embedding dimension (512 for CLIP, 384 for MiniLM)
    """
    # model.encode() takes a list of strings and returns a numpy array
    # show_progress_bar=True prints a tqdm progress bar during encoding
    embeddings = model.encode(texts, show_progress_bar=True, convert_to_numpy=True)
    return embeddings


def embed_chunks(model, chunks):
    """
    [NEW] Embed a mixed list of text strings and image dicts.

    Why a separate function? embed_texts() only handles strings. When our chunks
    list contains image dicts, we need to call model.encode() differently for each:
      - Text chunks:  model.encode(text_string)  → 512-dim vector
      - Image chunks: model.encode(PIL.Image)    → 512-dim vector
    CLIP handles both through the same encode() call, but the input type differs.

    Parameters:
        model  (SentenceTransformer): CLIP model
        chunks (list): mix of str and image dicts

    Returns:
        embeddings (np.ndarray): shape (N, 512) — one vector per chunk
    """
    inputs = []
    for chunk in chunks:
        if isinstance(chunk, dict) and chunk.get("type") == "image":
            # [NEW] For image chunks, pass a PIL Image object to the CLIP encoder
            # CLIP was trained to embed images and text into the same space,
            # so model.encode(PIL_image) returns a 512-dim vector comparable to text vectors.
            try:
                import io
                img_bytes = base64.b64decode(chunk["b64_data"])
                pil_img = Image.open(io.BytesIO(img_bytes))
                inputs.append(pil_img)
            except Exception as e:
                print(f"  [EMBED] Could not decode image {chunk.get('path')}: {e}")
                inputs.append("")  # fallback: empty string gives near-zero vector
        else:
            # Text chunk — pass as string (same as before)
            inputs.append(chunk if isinstance(chunk, str) else "")

    embeddings = model.encode(inputs, show_progress_bar=True, convert_to_numpy=True)
    return embeddings


# =============================================================================
# STEP 4 — VECTOR STORE (in-memory, using numpy)
# =============================================================================
#
# A vector store stores:
#   - the original text chunks (now also image dicts)
#   - their corresponding embedding vectors
#
# At search time, we:
#   1. Embed the user's query
#   2. Compare the query vector to all stored chunk vectors using cosine similarity
#   3. Return the top-K most similar chunks
#
# Cosine similarity formula:
#   similarity = (A · B) / (||A|| * ||B||)
# Where A and B are vectors, · is dot product, ||.|| is the L2 norm (magnitude).
# Result is between -1 and 1. For text, usually between 0 and 1.
#
# [NOTE] The vector store class itself is UNCHANGED — it works the same way
# regardless of whether chunks are text strings or image dicts. Only the
# embedding step above (embed_chunks) needed to change.

class SimpleVectorStore:
    """
    A minimal in-memory vector store using numpy for cosine similarity search.

    Attributes:
        chunks     (list): the text chunks (str) or image dicts
        embeddings (np.ndarray): shape (N, D) — one row per chunk
    """

    def __init__(self):
        self.chunks = []           # stores the text of each chunk (or image dict)
        self.embeddings = None     # stores all embedding vectors as a 2D numpy array

    def add(self, chunks, embeddings):
        """
        Add chunks and their precomputed embeddings to the store.

        Parameters:
            chunks     (list): text chunks or image dicts
            embeddings (np.ndarray): shape (N, D)
        """
        self.chunks = chunks
        self.embeddings = embeddings
        print(f"\n[VECTOR STORE] Indexed {len(chunks)} chunks.")
        print(f"[VECTOR STORE] Each embedding has {embeddings.shape[1]} dimensions.")

    def cosine_similarity(self, query_vec, doc_vecs):
        """
        Compute cosine similarity between one query vector and many document vectors.

        Parameters:
            query_vec (np.ndarray): shape (D,) — the query embedding
            doc_vecs  (np.ndarray): shape (N, D) — all chunk embeddings

        Returns:
            similarities (np.ndarray): shape (N,) — similarity score per chunk
        """
        dot_products = np.dot(doc_vecs, query_vec)
        query_norm = np.linalg.norm(query_vec)
        doc_norms = np.linalg.norm(doc_vecs, axis=1)
        doc_norms = np.maximum(doc_norms, 1e-10)
        similarities = dot_products / (query_norm * doc_norms)
        return similarities

    def search(self, query_embedding, top_k=3):
        """
        Retrieve the top-K most relevant chunks for a given query embedding.

        Parameters:
            query_embedding (np.ndarray): shape (D,) — the embedded user query
            top_k (int): number of top results to return

        Returns:
            results (list of dict): each dict has 'chunk' (str or image dict) and 'score' (float)
        """
        scores = self.cosine_similarity(query_embedding, self.embeddings)
        ranked_indices = np.argsort(scores)[::-1]
        top_indices = ranked_indices[:top_k]

        results = []
        for idx in top_indices:
            results.append({
                "chunk": self.chunks[idx],
                "score": float(scores[idx])
            })

        return results


# =============================================================================
# STEP 5 — RETRIEVER
# =============================================================================
#
# The retriever wraps the vector store and the embedding model together.
# Given a raw text query (string), it:
#   1. Embeds the query using the same CLIP model used to index documents
#   2. Searches the vector store for top-K similar chunks
# This is the "R" in RAG.
#
# [NOTE] This function is UNCHANGED — it receives a text query and returns
# top-K chunks (which may now be image dicts). The caller (generate_answer)
# handles the mixed content.

def retrieve(query, embedding_model, vector_store, top_k=3):
    """
    Retrieve the most relevant document chunks for a user query.

    Parameters:
        query           (str)               : the user's question as plain text
        embedding_model (SentenceTransformer): the loaded CLIP embedding model
        vector_store    (SimpleVectorStore) : the indexed vector store
        top_k           (int)               : how many chunks to retrieve

    Returns:
        results (list of dict): top-K chunks (text or image) with similarity scores
    """
    # Embed the query — same CLIP model used for indexing, so vectors are comparable
    query_embedding = embedding_model.encode(query, convert_to_numpy=True)

    # Search the vector store for the most similar chunks
    results = vector_store.search(query_embedding, top_k=top_k)

    return results


# =============================================================================
# STEP 6 — GENERATOR (Groq + LLaMA3-Vision)
# =============================================================================
#
# The generator takes:
#   - the retrieved context chunks (now possibly including images)
#   - the user's question
# and calls the Groq API to produce a final answer.
#
# [CHANGED] Model changed from 'llama-3.1-8b-instant' to 'meta-llama/llama-4-scout-17b-16e-instruct'.
#
# WHY THE NEW MODEL?
# ------------------
# 'llama-3.1-8b-instant' only understands text. If an image chunk is retrieved,
# the old model would have no way to reason about it.
# 'meta-llama/llama-4-scout-17b-16e-instruct' (LLaMA 3.2 11B Vision) is a multimodal model
# that can process base64-encoded images alongside text in the same API call.
# This is done using the OpenAI-style multi-content message format:
#
#   {"role": "user", "content": [
#       {"type": "text", "text": "Here is an image:"},
#       {"type": "image_url", "image_url": {"url": "data:image/jpeg;base64,<b64>"}},
#       {"type": "text", "text": "QUESTION: ..."},
#   ]}
#
# The old text-only message format is kept as a comment below for reference.

def generate_answer(question, retrieved_chunks, groq_client,
                    model="meta-llama/llama-4-scout-17b-16e-instruct"):
    """
    Generate an answer using the Groq vision-capable LLM, grounded in retrieved context.

    [CHANGED]
    - Default model changed from 'llama-3.1-8b-instant' to 'meta-llama/llama-4-scout-17b-16e-instruct'
    - User message now uses the multi-content format (list of content blocks) instead of
      a plain text string, so images can be included alongside text.

    Parameters:
        question         (str)          : the user's original question
        retrieved_chunks (list of dict) : output of retrieve() — chunks (text or image) + scores
        groq_client      (Groq)         : authenticated Groq API client
        model            (str)          : Groq model ID to use

    Returns:
        answer (str): the LLM's generated response
    """
    # --- System message (unchanged) ---
    system_message = (
        "You are a helpful AI assistant specializing in explaining RAG concepts. "
        "Answer the user's question ONLY based on the provided context below. "
        "If the context does not contain enough information to answer, say so clearly. "
        "Do not use any knowledge outside of what is in the context."
    )

    # --- Build the user message content as a list of content blocks ---
    # [CHANGED] Previously this was a single string:
    #   user_message = f"CONTEXT:\n...\nQUESTION: {question}"
    #
    # Now it is a list of dicts (content blocks) following the vision API format.
    # This lets us interleave text descriptions and base64 images freely.

    content_blocks = []

    # Opening context header (text block)
    content_blocks.append({
        "type": "text",
        "text": "CONTEXT (retrieved from knowledge base):\n" + "=" * 60
    })

    # Add each retrieved chunk as either a text block or an image block
    for i, result in enumerate(retrieved_chunks, start=1):
        chunk = result["chunk"]
        score = result["score"]

        if isinstance(chunk, dict) and chunk.get("type") == "image":
            # [NEW] Image chunk — add a text label then the image content block
            content_blocks.append({
                "type": "text",
                "text": f"\n[Image {i} | similarity: {score:.3f}] (from: {Path(chunk['path']).name})"
            })
            # The Groq vision API uses 'image_url' with a data URI for base64 images.
            # Format: "data:<mime_type>;base64,<base64_string>"
            content_blocks.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{chunk['mime_type']};base64,{chunk['b64_data']}"
                }
            })
        else:
            # Text chunk — same as before but now as a content block dict
            chunk_text = chunk if isinstance(chunk, str) else ""
            content_blocks.append({
                "type": "text",
                "text": f"\n[Chunk {i} | similarity: {score:.3f}]\n{chunk_text}"
            })

    # Closing separator and the user's question
    content_blocks.append({
        "type": "text",
        "text": "\n" + "=" * 60 + f"\n\nQUESTION: {question}"
    })

    # --- Call the Groq API ---
    # [CHANGED] 'content' is now a list of content blocks (supports images),
    # not a plain string. This is the key structural change from the old code.
    #
    # [OLD] Plain text message format — worked only for text:
    # response = groq_client.chat.completions.create(
    #     model="llama-3.1-8b-instant",
    #     messages=[
    #         {"role": "system", "content": system_message},
    #         {"role": "user",   "content": user_message},   ← plain string
    #     ],
    #     max_tokens=512,
    #     temperature=0.9,
    # )

    # [NEW] Multi-content message format — supports text + images:
    response = groq_client.chat.completions.create(
        model=model,                       # meta-llama/llama-4-scout-17b-16e-instruct
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user",   "content": content_blocks},  # ← list of blocks
        ],
        max_tokens=512,
        temperature=0.9,
    )

    answer = response.choices[0].message.content
    return answer


# =============================================================================
# STEP 7 — MAIN: Wire Everything Together
# =============================================================================
#
# This function runs the full RAG pipeline:
#   1. Load and chunk documents (now including images)
#   2. Load CLIP embedding model (multimodal)
#   3. Embed all chunks (text and images) and build the vector store
#   4. Set up the Groq LLM client (vision-capable model)
#   5. Enter an interactive Q&A loop in the terminal

def main():
    print("=" * 65)
    print("  RAG DEMO — Retrieval-Augmented Generation (Study Edition)")
    # [CHANGED] Updated banner to reflect new models
    print("  LLM: Groq / LLaMA-3.2-11B-Vision   |   Embeddings: CLIP-ViT-B-32")
    # [OLD] print("  LLM: Groq / LLaMA-3.1-8B-Instant   |   Embeddings: MiniLM-L6-v2")
    print("=" * 65)

    # -------------------------------------------------------------------------
    # 7a. Read the Groq API key
    # -------------------------------------------------------------------------
    api_key = os.environ.get("GROQ_API_KEY")

    if not api_key:
        env_file = Path(__file__).parent / ".env"

        if env_file.exists():
            for line in env_file.read_text().splitlines():
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" in line:
                    key, _, value = line.partition("=")
                    if key.strip() == "GROQ_API_KEY":
                        api_key = value.strip()
                        print("\n[API] Groq API key loaded from .env file.")
                        break

    if not api_key:
        print("\nERROR: GROQ_API_KEY not found.")
        print("Either:")
        print("  1. Create a .env file in this folder with the line:")
        print("       GROQ_API_KEY=your_key_here")
        print("  2. Or set the environment variable:")
        print("       PowerShell :  $env:GROQ_API_KEY = 'your_key_here'")
        print("       CMD        :  set GROQ_API_KEY=your_key_here")
        print("\nGet a free key at: https://console.groq.com")
        sys.exit(1)

    if not api_key.startswith("gsk_"):
        print(f"\nWARNING: Key doesn't look like a Groq key (expected 'gsk_...').")
        print("Continuing anyway — the API call will fail if the key is wrong.")
    else:
        print(f"\n[API] Groq API key found (ends with ...{api_key[-6:]}).")

    # -------------------------------------------------------------------------
    # 7b. Chunk the documents (text is split; images pass through as-is)
    # -------------------------------------------------------------------------
    print("\n[CHUNKING] Splitting documents into chunks...")
    chunks = chunk_documents(DOCUMENTS, chunk_size=500, chunk_overlap=50)
    text_chunks = [c for c in chunks if isinstance(c, str)]
    image_chunks = [c for c in chunks if isinstance(c, dict)]
    print(f"[CHUNKING] Created {len(text_chunks)} text chunks + {len(image_chunks)} image chunks.")

    if chunks:
        first_text = next((c for c in chunks if isinstance(c, str)), None)
        if first_text:
            print(f"\n[CHUNKING] Sample chunk (first text chunk):")
            print("-" * 40)
            print(textwrap.fill(first_text, width=60))
            print("-" * 40)

    # -------------------------------------------------------------------------
    # 7c. Load CLIP embedding model
    # [CHANGED] Now loads clip-ViT-B-32 instead of all-MiniLM-L6-v2
    # -------------------------------------------------------------------------
    embedding_model = load_embedding_model()  # defaults to clip-ViT-B-32
    # [OLD] embedding_model = load_embedding_model("all-MiniLM-L6-v2")

    # -------------------------------------------------------------------------
    # 7d. Embed all chunks (text and images) and build the vector store
    # [CHANGED] Now uses embed_chunks() instead of embed_texts() to handle images
    # -------------------------------------------------------------------------
    print("\n[INDEXING] Embedding all chunks — this is the INDEXING phase...")
    print("  (Each chunk is converted to a 512-dimensional CLIP vector.)")
    # [OLD] print("  (Each chunk is converted to a 384-dimensional vector.)")

    # [CHANGED] embed_chunks() handles mixed text + image input
    # [OLD] chunk_embeddings = embed_texts(embedding_model, chunks)
    chunk_embeddings = embed_chunks(embedding_model, chunks)

    vector_store = SimpleVectorStore()
    vector_store.add(chunks, chunk_embeddings)

    # -------------------------------------------------------------------------
    # 7e. Set up the Groq API client
    # -------------------------------------------------------------------------
    print("\n[LLM] Connecting to Groq API...")
    groq_client = Groq(api_key=api_key)
    print("[LLM] Groq client ready. Model: meta-llama/llama-4-scout-17b-16e-instruct")
    # [OLD] print("[LLM] Groq client ready. Model: llama-3.1-8b-instant")

    # -------------------------------------------------------------------------
    # 7f. Interactive Q&A loop
    # -------------------------------------------------------------------------
    print("\n" + "=" * 65)
    print("  RAG SYSTEM READY — Ask questions about your documents + images!")
    print("  Type 'quit' or 'exit' to stop.")
    print("=" * 65)

    print("\nSuggested questions to try:")
    print("  • What is RAG?")
    print("  • How does cosine similarity work?")
    print("  • What is chunking and why is it needed?")
    print("  • What is an embedding?")
    print("  • Describe what you see in the uploaded image.")

    while True:
        print()
        try:
            user_question = input("Your question: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n\n[EXIT] Interrupted. Goodbye!")
            break

        if user_question.lower() in ("quit", "exit", "q", ""):
            print("\n[EXIT] Goodbye! Thanks for studying RAG.")
            break

        print(f"\n[RETRIEVAL] Searching for relevant chunks for: '{user_question}'")

        results = retrieve(
            query=user_question,
            embedding_model=embedding_model,
            vector_store=vector_store,
            top_k=3
        )

        print("\n[RETRIEVAL] Top-3 chunks retrieved:")
        for i, r in enumerate(results, start=1):
            chunk = r["chunk"]
            if isinstance(chunk, dict) and chunk.get("type") == "image":
                print(f"\n  Chunk {i} (similarity score: {r['score']:.4f}): [IMAGE] {chunk['path']}")
            else:
                print(f"\n  Chunk {i} (similarity score: {r['score']:.4f}):")
                print("  " + "-" * 50)
                wrapped = textwrap.fill(chunk, width=60, initial_indent="  ", subsequent_indent="  ")
                print(wrapped)
                print("  " + "-" * 50)

        print("\n[GENERATION] Sending context + question to Groq (LLaMA-3.2-11B-Vision)...")
        # [OLD] print("\n[GENERATION] Sending context + question to Groq (LLaMA-3.1-8B-Instant)...")
        try:
            answer = generate_answer(
                question=user_question,
                retrieved_chunks=results,
                groq_client=groq_client,
                model="meta-llama/llama-4-scout-17b-16e-instruct"
                # [OLD] model="llama-3.1-8b-instant"
            )
        except Exception as e:
            print(f"\n[ERROR] Groq API call failed: {e}")
            continue

        print("\n" + "=" * 65)
        print("  ANSWER (generated by LLaMA3.2-11B-Vision via Groq):")
        # [OLD] print("  ANSWER (generated by LLaMA3-8B via Groq):")
        print("=" * 65)
        print(textwrap.fill(answer, width=65))
        print("=" * 65)


# =============================================================================
# ENTRY POINT
# =============================================================================
if __name__ == "__main__":
    main()
